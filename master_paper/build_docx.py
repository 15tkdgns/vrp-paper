# -*- coding: utf-8 -*-
"""
학위논문 docx 생성기 (완전판)
- 위첨자 자동 처리 (¹²³ → Word superscript)
- 표 자동 변환 (| 구분자 → Word 표 객체)
- 참고문헌 저널명 이탤릭
- 페이지 번호 (하단 가운데)
- 장별 페이지 나누기
- 양식.txt 여백·폰트 적용
"""

import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from copy import deepcopy

# ── 경로 설정 ──────────────────────────────────────────────────────────────
BASE      = Path(r"C:\Users\user\Desktop\vrp\master_paper")
FILES     = ["front.txt","ch1.txt","ch2.txt","ch3.txt","ch4.txt","ch5.txt","references.txt","appendix.txt"]
OUT       = BASE / "0421_thesis.docx"
SRC_FRONT = BASE / "국문 MS word 학위논문 서식.docx"

FONT_KO   = "HY견명조"
FONT_EN   = "Times New Roman"
BODY_SIZE = 11

# 위첨자 유니코드 매핑
SUP_MAP = {'⁰':'0','¹':'1','²':'2','³':'3','⁴':'4',
           '⁵':'5','⁶':'6','⁷':'7','⁸':'8','⁹':'9'}
SUP_CHARS = set(SUP_MAP.keys())

# 아래첨자 유니코드 매핑
SUB_MAP = {'₀':'0','₁':'1','₂':'2','₃':'3','₄':'4',
           '₅':'5','₆':'6','₇':'7','₈':'8','₉':'9'}
SUB_CHARS = set(SUB_MAP.keys())

# 단일 소문자 _x 아래첨자 자동 처리 여부 (appendix 의사코드에서는 False)
_bare_sub = True

# ── 문서 초기화 ────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_height   = Mm(297)
sec.page_width    = Mm(210)
sec.top_margin    = Mm(35)
sec.bottom_margin = Mm(25)
sec.left_margin   = Mm(35)
sec.right_margin  = Mm(30)

def set_font(run, size=10, bold=False, italic=False, superscript=False, subscript=False):
    run.font.name   = FONT_EN
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    # 한글 폰트
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_KO)
    if not rPr.findall(qn('w:rFonts')):
        rPr.insert(0, rFonts)
    # 위첨자 / 아래첨자
    if superscript or subscript:
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript' if superscript else 'subscript')
        rPr.append(vertAlign)

def set_para_fmt(para, size=BODY_SIZE, bold=False,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 sp_before=0, sp_after=6,
                 ls=2.0, indent_first=0):
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before       = Pt(sp_before)
    pf.space_after        = Pt(sp_after)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = ls
    if indent_first:
        pf.first_line_indent = Pt(indent_first)

def add_runs_with_sup(para, text, size=10, bold=False, italic=False):
    """위첨자·아래첨자 포함 텍스트를 run 분리해서 추가
    처리: 유니코드 위첨자(¹²³…), 유니코드 아래첨자(₀₁₂…),
          _{...} 아래첨자, _(...) 아래첨자,
          ^{...} 위첨자, ^word 위첨자 (^k, ^22, ^hat 등)
    """
    def flush(s, sup=False, sub=False):
        if s:
            r = para.add_run(s)
            set_font(r, size=size, bold=bold, italic=italic,
                     superscript=sup, subscript=sub)

    def extract_braced(text, i):
        """i는 '{' 다음 위치. 대응 '}' 까지 내용 반환 (i_after, content)."""
        depth, s = 1, ""
        while i < len(text):
            ch = text[i]; i += 1
            if ch == '{':
                depth += 1; s += ch
            elif ch == '}':
                depth -= 1
                if depth:
                    s += ch
                else:
                    break
            else:
                s += ch
        return i, s

    def extract_paren(text, i):
        """i는 '(' 다음 위치. 대응 ')' 까지 내용 반환 (i_after, content)."""
        depth, s = 1, ""
        while i < len(text):
            ch = text[i]; i += 1
            if ch == '(':
                depth += 1; s += ch
            elif ch == ')':
                depth -= 1
                if depth:
                    s += ch
                else:
                    break
            else:
                s += ch
        return i, s

    buf = ""
    i = 0
    n = len(text)

    while i < n:
        c = text[i]

        # ── 유니코드 위첨자 (¹²³…) ──────────────────────────────────
        if c in SUP_CHARS:
            flush(buf); buf = ""
            s = ""
            while i < n and text[i] in SUP_CHARS:
                s += SUP_MAP[text[i]]; i += 1
            flush(s, sup=True)
            continue

        # ── 유니코드 아래첨자 (₀₁₂…) ────────────────────────────────
        if c in SUB_CHARS:
            flush(buf); buf = ""
            s = ""
            while i < n and text[i] in SUB_CHARS:
                s += SUB_MAP[text[i]]; i += 1
            flush(s, sub=True)
            continue

        # ── _{...} 아래첨자 ──────────────────────────────────────────
        if c == '_' and i + 1 < n and text[i + 1] == '{':
            flush(buf); buf = ""
            i, s = extract_braced(text, i + 2)
            flush(s, sub=True)
            continue

        # ── _(...) 아래첨자 ──────────────────────────────────────────
        if c == '_' and i + 1 < n and text[i + 1] == '(':
            flush(buf); buf = ""
            i, s = extract_paren(text, i + 2)
            flush(s, sub=True)
            continue

        # ── ^{...} 위첨자 ────────────────────────────────────────────
        if c == '^' and i + 1 < n and text[i + 1] == '{':
            flush(buf); buf = ""
            i, s = extract_braced(text, i + 2)
            flush(s, sup=True)
            continue

        # ── ^word 위첨자 (^{ 는 위에서 처리, ^ 뒤 알파뉴메릭 연속) ──
        if c == '^' and i + 1 < n and text[i + 1].isalnum():
            flush(buf); buf = ""
            i += 1
            s = ""
            while i < n and text[i].isalnum():
                s += text[i]; i += 1
            flush(s, sup=True)
            continue

        # ── _x 단일 소문자/숫자 아래첨자 (appendix 의사코드 제외) ───
        if _bare_sub and c == '_' and i + 1 < n:
            nc = text[i + 1]
            if (nc.islower() or nc.isdigit()) and (i + 2 >= n or not text[i + 2].isalnum()):
                flush(buf); buf = ""
                flush(nc, sub=True)
                i += 2
                continue

        buf += c
        i += 1

    flush(buf)

# ── 페이지 번호 (하단 가운데) ──────────────────────────────────────────────
def add_page_number(section):
    footer = section.footer
    para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    set_font(run, size=10)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2])

add_page_number(sec)

# ── 줄 분류 ──────────────────────────────────────────────────────────────
CH_PAT    = re.compile(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅺⅻ]\.')
SEC_PAT   = re.compile(r'^[1-9][0-9]?\.\s')
SUB_PAT   = re.compile(r'^[1-9][0-9]?\)\s')
TCAP_PAT  = re.compile(r'^(표\s+[0-9A-Z][-0-9]*\.|Table\s+[0-9A-Z][-0-9]*\.)', re.I)
FCAP_PAT  = re.compile(r'^(그림\s+[0-9]+\.|Fig\.\s+[0-9]+\.)', re.I)
PIPE_PAT  = re.compile(r'\|')
SEP_PAT   = re.compile(r'^-{3,}$')
REF_PAT   = re.compile(r'^\s*[0-9]+\.\s+\w')

def classify(line):
    s = line.strip()
    if not s:                    return 'blank'
    if SEP_PAT.match(s):         return 'sep'
    if CH_PAT.match(s):          return 'ch'
    if SEC_PAT.match(s):         return 'sec'
    if SUB_PAT.match(s):         return 'sub'
    if TCAP_PAT.match(s):        return 'tcap'
    if FCAP_PAT.match(s):        return 'fcap'
    return 'body'

# ── 표 렌더링 (| 구분자) ──────────────────────────────────────────────────
def make_pipe_table(doc, lines):
    rows = []
    for ln in lines:
        if '|' in ln:
            cells = [c.strip() for c in ln.split('|')]
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols: break
            cell = tbl.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            add_runs_with_sup(para, cell_text, size=9, bold=(ri==0))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after  = Pt(1)
    doc.add_paragraph()  # 표 뒤 여백

# ── 참고문헌 줄 처리 (저널명 이탤릭) ────────────────────────────────────
def add_reference_line(doc, text):
    para = doc.add_paragraph()
    set_para_fmt(para, size=10, sp_before=2, sp_after=2, ls=1.3, indent_first=-20)
    para.paragraph_format.left_indent = Pt(20)

    # 패턴: "N. Author(s). YEAR. Title. Journal, Vol: Pages."
    # 저널명 = 제목(마침표) 뒤, 콤마 앞
    m = re.match(
        r'^(\s*[0-9]+\.\s+)'        # 번호
        r'(.+?\.\s+[0-9]{4}\.\s+)'  # 저자 + 연도
        r'(.+?\.\s+)'               # 제목
        r'([^,0-9]+)'               # 저널명
        r'(.*)',                     # 나머지
        text
    )
    if m:
        for i, (t, ital) in enumerate([
            (m.group(1), False),
            (m.group(2), False),
            (m.group(3), False),
            (m.group(4).strip(), True),   # 저널명 이탤릭
            (m.group(5), False),
        ]):
            if t:
                add_runs_with_sup(para, t, size=10, italic=ital)
    else:
        add_runs_with_sup(para, text, size=10)

# ── 캡션 추가 ────────────────────────────────────────────────────────────
def add_caption(doc, text):
    para = doc.add_paragraph()
    add_runs_with_sup(para, text, size=BODY_SIZE, bold=True)
    set_para_fmt(para, size=BODY_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER,
                 sp_before=3, sp_after=3, ls=2.0)

# ── 제목 추가 ────────────────────────────────────────────────────────────
def add_heading(doc, text, level):
    sizes   = {1:16, 2:14, 3:12}
    befores = {1:24, 2:12, 3:8}
    afters  = {1:12, 2:6,  3:4}
    lss     = {1:1.8, 2:1.5, 3:1.5}
    aligns  = {1:WD_ALIGN_PARAGRAPH.CENTER, 2:WD_ALIGN_PARAGRAPH.LEFT, 3:WD_ALIGN_PARAGRAPH.LEFT}
    para = doc.add_paragraph()
    add_runs_with_sup(para, text, size=sizes.get(level, BODY_SIZE), bold=True)
    set_para_fmt(para, size=sizes.get(level, BODY_SIZE), bold=True,
                 align=aligns.get(level, WD_ALIGN_PARAGRAPH.LEFT),
                 sp_before=befores.get(level, 6),
                 sp_after=afters.get(level, 4),
                 ls=lss.get(level, 2.0))

# ── 본문 단락 ─────────────────────────────────────────────────────────────
def add_body(doc, text, indent=True, size=BODY_SIZE, italic=False):
    para = doc.add_paragraph()
    add_runs_with_sup(para, text, size=size, italic=italic)
    set_para_fmt(para, size=size, sp_before=0, sp_after=6,
                 indent_first=20 if indent else 0)

# ── 서식.docx 1~3페이지(표지/인준지) 앞에 복사 ───────────────────────────
def prepend_front_pages(doc, src_path):
    """서식.docx에서 '차례' 섹션 이전(표지+인준지)을 doc 앞에 삽입"""
    src  = Document(str(src_path))
    body = doc.element.body
    insert_idx = 0
    for elem in src.element.body:
        texts = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
        cleaned = texts.replace(' ', '').replace('　', '')
        if '차례' in cleaned or '목차' in cleaned:
            break
        body.insert(insert_idx, deepcopy(elem))
        insert_idx += 1
    # 4페이지 시작용 명시적 페이지 브레이크
    pg_p = OxmlElement('w:p')
    pg_r = OxmlElement('w:r')
    pg_br = OxmlElement('w:br')
    pg_br.set(qn('w:type'), 'page')
    pg_r.append(pg_br)
    pg_p.append(pg_r)
    body.insert(insert_idx, pg_p)
    print(f"  표지/인준지 {insert_idx}개 elements 복사 완료")

prepend_front_pages(doc, SRC_FRONT)

# ── 차례 삽입 (front.txt 처리 전) ────────────────────────────────────────
from build_toc import build_toc_section
build_toc_section(doc)

# ── 메인 처리 ─────────────────────────────────────────────────────────────
in_reference = False
in_appendix  = False
skip_toc     = False   # 목차/표차례/그림차례 텍스트 건너뜀

for fname in FILES:
    fpath = BASE / fname
    if not fpath.exists():
        print(f"  skip: {fname}")
        continue

    lines = fpath.read_text(encoding='utf-8').splitlines()
    print(f"  {fname}: {len(lines)} lines")

    if fname == 'appendix.txt':
        in_appendix = True
        _bare_sub = False  # 의사코드 변수명 오탐 방지

    i = 0
    while i < len(lines):
        line  = lines[i]
        s     = line.strip()
        kind  = classify(line)

        if kind == 'blank':
            i += 1
            continue

        # 구분선 → 페이지 나누기 / 목차 skip 해제
        if kind == 'sep':
            skip_toc = False
            doc.add_page_break()
            i += 1
            continue

        # 목차/표차례/그림차례 텍스트 건너뜀
        if fname == 'front.txt' and s in ('목차', '표 차례', '그림 차례'):
            skip_toc = True
            i += 1
            continue
        if skip_toc:
            i += 1
            continue

        # 참고문헌 섹션 감지
        if s == '참고문헌':
            doc.add_page_break()
            in_reference = True
            para = doc.add_paragraph()
            add_runs_with_sup(para, '참고문헌', size=14, bold=True)
            set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                         sp_before=0, sp_after=12)
            i += 1
            continue

        # 참고문헌 항목 (sec/sub 패턴보다 먼저 처리)
        if in_reference and REF_PAT.match(s):
            add_reference_line(doc, s)
            i += 1
            continue

        # 장 제목
        if kind == 'ch':
            if fname != 'front.txt':
                doc.add_page_break()
            in_reference = False
            add_heading(doc, s, 1)
            i += 1
            continue

        # 절 제목
        if kind == 'sec':
            add_heading(doc, s, 2)
            i += 1
            continue

        # 소절 제목
        if kind == 'sub':
            add_heading(doc, s, 3)
            i += 1
            continue

        # 표 캡션
        if kind == 'tcap':
            cap_line = s
            i += 1
            # 영문 캡션 있으면 합치기
            if i < len(lines) and classify(lines[i]) == 'tcap':
                cap_line += '  /  ' + lines[i].strip()
                i += 1
            add_caption(doc, cap_line)

            # 표 데이터 수집 (초기 빈 줄 건너뜀)
            tbl_lines = []
            # 빈 줄 skip
            while i < len(lines) and not lines[i].strip():
                i += 1
            # 데이터 수집
            blank_run = 0
            while i < len(lines):
                nxt = lines[i].strip()
                if classify(lines[i]) in ('ch','sec','sub','tcap','fcap','sep'):
                    break
                if not nxt:
                    blank_run += 1
                    if blank_run >= 1:  # 빈 줄 1개면 표 종료 (본문 흡수 방지)
                        break
                    i += 1
                    continue
                blank_run = 0
                tbl_lines.append(nxt)
                i += 1

            if tbl_lines:
                has_pipe = any('|' in l for l in tbl_lines)
                if has_pipe:
                    make_pipe_table(doc, tbl_lines)
                else:
                    for tl in tbl_lines:
                        para = doc.add_paragraph()
                        add_runs_with_sup(para, tl, size=9)
                        set_para_fmt(para, size=9, sp_before=0,
                                     sp_after=1, ls=1.2)
            continue

        # 그림 캡션
        if kind == 'fcap':
            cap_line = s
            i += 1
            if i < len(lines) and classify(lines[i]) == 'fcap':
                cap_line += '  /  ' + lines[i].strip()
                i += 1
            # 이미지 위치 표시
            ph = doc.add_paragraph()
            ph_run = ph.add_run('[이미지 삽입 위치]')
            set_font(ph_run, size=10, italic=True)
            set_para_fmt(ph, align=WD_ALIGN_PARAGRAPH.CENTER,
                         sp_before=6, sp_after=3, ls=1.2)
            add_caption(doc, cap_line)
            continue


        # 부록 제목
        if in_appendix and s.startswith('부록 '):
            add_heading(doc, s, 2)
            i += 1
            continue

        # 일반 본문
        add_body(doc, s, indent=not in_reference and not in_appendix)
        i += 1

# ── 저장 ──────────────────────────────────────────────────────────────────
doc.save(str(OUT))
sz = OUT.stat().st_size
print(f"\n[완료] {OUT.name}  ({sz:,} bytes / {sz//1024} KB)")
