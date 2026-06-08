# -*- coding: utf-8 -*-
"""
차례(목차) + 표차례 + 그림차례 생성
참고논문.pdf 레이아웃 기준:
  장(Ⅰ): 들여쓰기 없음, 점선 없음, 굵게, 우측 페이지번호
  절(1.): 7mm 들여쓰기, 점선, 우측 페이지번호
  소절(1)): 14mm 들여쓰기, 점선, 우측 페이지번호
  표차례/그림차례: 절과 같은 형식
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re

BASE = Path(r"C:\Users\user\Desktop\vrp\master_paper")
FRONT = BASE / "front.txt"

FONT_KO = "신명조"
FONT_EN = "Times New Roman"

# ── 텍스트 너비(mm): 210 - 35 - 30 = 145mm → 탭 위치 ──────────────────
TEXT_WIDTH_MM = 145.0  # mm
TAB_POS_TWIPS = int(TEXT_WIDTH_MM * 56.69)  # 1mm = 56.69 twips

# ── 들여쓰기(참고논문 좌표 기준) ──────────────────────────────────────────
INDENT = {
    1: Mm(0),    # 장 (Ⅰ, Ⅱ)
    2: Mm(7),    # 절 (1., 2.)
    3: Mm(14),   # 소절 (1), 2))
    'tbl': Mm(0),  # 표차례
    'fig': Mm(0),  # 그림차례
}

def set_font(run, size=10.5, bold=False, italic=False):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_KO)
    rFonts.set(qn('w:ascii'), FONT_EN)

def add_right_tab(para, with_dot=True):
    """오른쪽 정렬 탭 + 점선 리더 추가"""
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None:
        pPr.remove(tabs)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(TAB_POS_TWIPS))
    tab.set(qn('w:leader'), 'dot' if with_dot else 'none')
    tabs.append(tab)
    pPr.append(tabs)

def set_para_spacing(para, sp_before=0, sp_after=3, ls=1.3):
    pf = para.paragraph_format
    pf.space_before = Pt(sp_before)
    pf.space_after  = Pt(sp_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = ls

def add_toc_entry(doc, text, level, page_num=""):
    """
    차례 항목 추가
    level=1: 장 (굵게, 점선 없음)
    level=2: 절 (점선 있음)
    level=3: 소절 (더 들여쓰기, 점선 있음)
    """
    para = doc.add_paragraph()
    is_ch = (level == 1)
    size  = 11 if is_ch else 10.5

    # 들여쓰기
    pf = para.paragraph_format
    pf.left_indent = INDENT.get(level, Mm(0))
    set_para_spacing(para, sp_before=2 if is_ch else 0, sp_after=2 if is_ch else 1)

    # 탭 (장은 점선 없음, 절/소절은 점선)
    add_right_tab(para, with_dot=not is_ch)

    # 텍스트 run
    r1 = para.add_run(text)
    set_font(r1, size=size, bold=is_ch)

    # 탭 + 페이지번호 run
    r2 = para.add_run('\t' + page_num)
    set_font(r2, size=size, bold=is_ch)

    return para

def add_section_title(doc, text):
    """차례/표차례/그림차례 섹션 제목"""
    para = doc.add_paragraph()
    run  = para.add_run(text)
    set_font(run, size=14, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, sp_before=0, sp_after=18, ls=1.5)
    return para

def add_blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        set_para_spacing(p, sp_before=0, sp_after=0)

# ── 목차 파싱 ─────────────────────────────────────────────────────────────
def parse_toc_from_front(path):
    """front.txt의 목차 섹션 파싱"""
    text = path.read_text(encoding='utf-8')
    # 목차 ~ 표 차례 사이 추출
    m = re.search(r'목차\s*\n(.*?)(?=표 차례|그림 차례|---|\Z)',
                  text, re.DOTALL)
    if not m:
        return []

    entries = []
    CH_PAT   = re.compile(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅺⅻ]\.')
    SEC_PAT  = re.compile(r'^\s{1,3}[0-9]+\.\s')
    SUB_PAT  = re.compile(r'^\s{4,}[0-9]+\)\s')
    SKIP_PAT = re.compile(r'^(참고문헌|부록)')

    for line in m.group(1).splitlines():
        s = line.rstrip()
        if not s.strip(): continue
        if SKIP_PAT.match(s.strip()):
            entries.append((0, s.strip()))
            continue
        if CH_PAT.match(s.strip()):
            entries.append((1, s.strip()))
        elif SUB_PAT.match(s):
            entries.append((3, s.strip()))
        elif SEC_PAT.match(s):
            entries.append((2, s.strip()))
        else:
            entries.append((0, s.strip()))
    return entries

def parse_table_list(path):
    """표차례 항목 파싱"""
    text = path.read_text(encoding='utf-8')
    m = re.search(r'표 차례\s*\n(.*?)(?=그림 차례|---|\Z)', text, re.DOTALL)
    if not m: return []
    return [l.strip() for l in m.group(1).splitlines() if l.strip()]

def parse_fig_list(path):
    """그림차례 항목 파싱"""
    text = path.read_text(encoding='utf-8')
    m = re.search(r'그림 차례\s*\n(.*?)(?=초록|Abstract|---|\Z)', text, re.DOTALL)
    if not m: return []
    return [l.strip() for l in m.group(1).splitlines() if l.strip()]

# ── 메인: 차례 문서 생성 ──────────────────────────────────────────────────
def build_toc_section(doc):
    """기존 doc에 차례/표차례/그림차례 삽입"""

    # ── 1. 차례 ──────────────────────────────────────────────────────────
    add_section_title(doc, '차   례')
    toc_entries = parse_toc_from_front(FRONT)

    for level, text in toc_entries:
        if level == 0:
            # 참고문헌, 부록 등
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_font(r, size=11, bold=True)
            add_right_tab(p, with_dot=False)
            p.add_run('\t')
            set_para_spacing(p, sp_before=2, sp_after=2)
        else:
            add_toc_entry(doc, text, level, page_num='')

    doc.add_page_break()

    # ── 2. 표차례 ────────────────────────────────────────────────────────
    add_section_title(doc, '표   차   례')
    for item in parse_table_list(FRONT):
        para = doc.add_paragraph()
        pf   = para.paragraph_format
        pf.left_indent = Mm(0)
        set_para_spacing(para, sp_before=0, sp_after=3, ls=1.3)
        add_right_tab(para, with_dot=True)
        r = para.add_run(item + '\t')
        set_font(r, size=10.5)

    doc.add_page_break()

    # ── 3. 그림차례 ──────────────────────────────────────────────────────
    add_section_title(doc, '그   림   차   례')
    for item in parse_fig_list(FRONT):
        para = doc.add_paragraph()
        set_para_spacing(para, sp_before=0, sp_after=3, ls=1.3)
        add_right_tab(para, with_dot=True)
        r = para.add_run(item + '\t')
        set_font(r, size=10.5)

    doc.add_page_break()


# ── 독립 실행 테스트 ──────────────────────────────────────────────────────
if __name__ == '__main__':
    doc = Document()
    sec = doc.sections[0]
    from docx.shared import Mm as M
    sec.page_height   = M(297)
    sec.page_width    = M(210)
    sec.top_margin    = M(35)
    sec.bottom_margin = M(30)
    sec.left_margin   = M(35)
    sec.right_margin  = M(30)

    build_toc_section(doc)
    out = BASE / "toc_test.docx"
    doc.save(str(out))
    print(f"[완료] {out.name}  ({out.stat().st_size:,} bytes)")
